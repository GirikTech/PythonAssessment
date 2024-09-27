import customtkinter as ctk
from tkinter import messagebox
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from customtkinter import CTkTabview


# Change appearance mode to light
ctk.set_appearance_mode("Light")
# Set a colorful theme
ctk.set_default_color_theme("blue")

# Define User class to store user information
class User:
    def __init__(self, username, password, role, business):
        self.username = username
        self.password = password
        self.role = role
        self.business = business
# Define Product class to store product information
class Product:
    def __init__(self, name, price, added_by):
        self.name = name
        self.price = price
        self.added_by = added_by
# Main InvoiceMaker class
class InvoiceMaker:
    # Initialise the main window
    def __init__(self):
        self.root = ctk.CTk()
        self.root.title("Invoice Maker - Iteration 4")
        
        self.root.geometry("700x1000")
         # Initialise user and data variables
        self.current_user = None
        self.users = self.load_data("users.json")
        self.products = self.load_data("products.json")
        self.purchase_history = self.load_data("purchase_history.json")

        # Set up the user interface
        self.setup_ui()

    # Method to load data from JSON files
    def load_data(self, filename):
        try:
            with open(filename, 'r') as f:
                return json.load(f)
        except FileNotFoundError:
            return {}
        
    # Method to save data to JSON files
    def save_data(self, data, filename):
        with open(filename, 'w') as f:
            json.dump(data, f)

    # Method to set up the main user interface
    def setup_ui(self):
        # Create main frame
        self.main_frame = ctk.CTkFrame(self.root)
        self.main_frame.pack(expand=True, fill="both", padx=20, pady=20)

        # Create frames for different sections
        self.login_frame = ctk.CTkFrame(self.main_frame)
        self.signup_frame = ctk.CTkFrame(self.main_frame)
        self.product_frame = ctk.CTkFrame(self.main_frame)
        self.invoice_frame = ctk.CTkFrame(self.main_frame)
        self.history_frame = ctk.CTkFrame(self.main_frame)

        # Set up individual frames
        self.setup_login_frame()
        self.setup_signup_frame()
        self.setup_product_frame()
        self.setup_invoice_frame()
        self.setup_history_frame()

        # Show the login frame initially
        self.show_frame(self.login_frame)

    # Method to set up the login frame
    def setup_login_frame(self):
        # Add logo label
        self.logo_label = ctk.CTkLabel(self.login_frame, text="INVOICE WIZARD", font=("Arial Rounded MT Bold", 30))
        self.logo_label.pack(pady=(20, 30))

        # Add username and password entry fields
        ctk.CTkLabel(self.login_frame, text="Username:", font=("Arial Rounded MT Bold", 14)).pack(pady=5)
        self.username_entry = ctk.CTkEntry(self.login_frame, width=300)
        self.username_entry.pack(pady=5)
        ctk.CTkLabel(self.login_frame, text="Password:", font=("Arial Rounded MT Bold", 14)).pack(pady=5)
        self.password_entry = ctk.CTkEntry(self.login_frame, show="*", width=300)
        self.password_entry.pack(pady=5)
        
        # Add login and signup buttons
        ctk.CTkButton(self.login_frame, text="Login", command=self.login, fg_color="#4CAF50", hover_color="#45a049", width=200).pack(pady=10)
        ctk.CTkButton(self.login_frame, text="Sign Up", command=lambda: self.show_frame(self.signup_frame), fg_color="#2196F3", hover_color="#1976D2", width=200).pack(pady=5)

    # Method to set up the signup frame
    def setup_signup_frame(self):
        ctk.CTkLabel(self.signup_frame, text="Create a Account", font=("Arial Rounded MT Bold", 24, "bold")).pack(pady=(20, 30))
        
        # Add entry fields for new account
        ctk.CTkLabel(self.signup_frame, text="Username:", font=("Arial Rounded MT Bold", 14)).pack(pady=5)
        self.new_username_entry = ctk.CTkEntry(self.signup_frame, width=300)
        self.new_username_entry.pack(pady=5)
        ctk.CTkLabel(self.signup_frame, text="Password:", font=("Arial Rounded MT Bold", 14)).pack(pady=5)
        self.new_password_entry = ctk.CTkEntry(self.signup_frame, show="*", width=300)
        self.new_password_entry.pack(pady=5)
        ctk.CTkLabel(self.signup_frame, text="Business:", font=("Arial Rounded MT Bold", 14)).pack(pady=5)
        self.business_entry = ctk.CTkEntry(self.signup_frame, width=300)
        self.business_entry.pack(pady=5)
        
        # Add signup and back buttons
        ctk.CTkButton(self.signup_frame, text="Sign Up", command=self.signup, fg_color="#4CAF50", hover_color="#45a049", width=200).pack(pady=10)
        ctk.CTkButton(self.signup_frame, text="Back to Login", command=lambda: self.show_frame(self.login_frame), fg_color="#2196F3", hover_color="#1976D2", width=200).pack(pady=5)

    # Method to set up the product frame
    def setup_product_frame(self):
        # Create a tab view
        self.tab_view = ctk.CTkTabview(self.product_frame)
        self.tab_view.pack(fill="both", expand=True, padx=20, pady=20)

        # Create tabs
        self.invoice_tab = self.tab_view.add("Invoice")
        self.store_tab = self.tab_view.add("Store")

        # Set up Invoice tab
        self.setup_invoice_tab()

        # Set up Store tab (only visible for owners)
        self.setup_store_tab()

    # Method to set up the invoice tab
    def setup_invoice_tab(self):
    # Product selection and invoice generation (for all users)
        ctk.CTkLabel(self.invoice_tab, text="Select Product", font=("Arial Rounded MT Bold", 18, "bold")).pack(pady=(20, 10))
        
        # Add product selection dropdown
        self.product_listbox = ctk.CTkOptionMenu(self.invoice_tab, values=[], command=self.update_price_display, width=300)
        self.product_listbox.pack(pady=10)
        
        # Add price display label
        self.price_display = ctk.CTkLabel(self.invoice_tab, text="Price: $0.00", font=("Arial", 16))
        self.price_display.pack(pady=5)
        
        # Add quantity entry field
        ctk.CTkLabel(self.invoice_tab, text="Quantity:", font=("Arial", 14)).pack(pady=(20,10))
        self.quantity_entry = ctk.CTkEntry(self.invoice_tab, width=200)
        self.quantity_entry.pack(pady=5)
        
        self.selected_products = []
        
        # Add button to add product to invoice
        ctk.CTkButton(self.invoice_tab, text="Add to Invoice", command=self.add_to_invoice, 
                    fg_color="#4CAF50", hover_color="#45a049", width=250).pack(pady=20, )
        
        # Add text box to display selected products
        self.selected_products_text = ctk.CTkTextbox(self.invoice_tab, width=550, height=250, text_color="black")
        self.selected_products_text.pack(pady=20)
        
        # Add buttons for generating invoice, viewing history, and logout
        button_frame = ctk.CTkFrame(self.invoice_tab)
        button_frame.pack(pady=20)
        ctk.CTkButton(button_frame, text="Generate Invoice", command=self.show_invoice, 
                    fg_color="#2196F3", hover_color="#1976D2", width=200).pack(side='left', padx=10)
        ctk.CTkButton(button_frame, text="View Purchase History", command=self.show_history, 
                    fg_color="#FF9800", hover_color="#45a049", width=200).pack(side='left', padx=10)
        ctk.CTkButton(button_frame, text="Logout", command=self.logout, 
                    fg_color="#f44336", hover_color="#d32f2f", width=200).pack(side='left', padx=10)

     # Method to set up the store tab
    def setup_store_tab(self):
        # Add new product section
        ctk.CTkLabel(self.store_tab, text="Add New Product", font=("Arial Rounded MT Bold", 22, "bold")).pack(pady=(15,5))
        
        ctk.CTkLabel(self.store_tab, text="Product Name:", font=("Arial", 14)).pack(pady=(20,5))
        self.product_name_entry = ctk.CTkEntry(self.store_tab, width=400)
        self.product_name_entry.pack(pady=5)
        
        ctk.CTkLabel(self.store_tab, text="Product Price:", font=("Arial", 14)).pack(pady=(15,5))
        self.product_price_entry = ctk.CTkEntry(self.store_tab, width=400)
        self.product_price_entry.pack(pady=5)
        
        ctk.CTkButton(self.store_tab, text="Add Product", command=self.add_product, 
                    fg_color="#4CAF50", hover_color="#45a049", width=250).pack(pady=20)

        # Add employee management section
        ctk.CTkLabel(self.store_tab, text="Employee Management", font=("Arial Rounded MT Bold", 22, "bold")).pack(pady=(35,15))
        ctk.CTkButton(self.store_tab, text="Create Employee Account", command=self.create_employee_account, 
                    fg_color="#4CAF50", hover_color="#45a049", width=300).pack(pady=10)
        ctk.CTkButton(self.store_tab, text="Change Employee Password", command=self.change_employee_password, 
                    fg_color="#FF9800", hover_color="#F57C00", width=300).pack(pady=10)

        # Add bank account section
        ctk.CTkLabel(self.store_tab, text="Bank Account", font=("Arial Rounded MT Bold", 22, "bold")).pack(pady=(40,20))
        ctk.CTkLabel(self.store_tab, text="Account Number:", font=("Arial", 14)).pack(pady=(15,5))
        self.bank_account_entry = ctk.CTkEntry(self.store_tab, width=400)
        self.bank_account_entry.pack(pady=5)
        ctk.CTkButton(self.store_tab, text="Update Bank Account", command=self.update_bank_account, 
                    fg_color="#4CAF50", hover_color="#45a049", width=250).pack(pady=20)
    
    # Method to update bank account
    def update_bank_account(self):
        account_number = self.bank_account_entry.get().strip()
        if account_number:
            if 'bank_account' in self.users[self.current_user.username]:
                self.users[self.current_user.username]['bank_account'] = account_number
                self.save_data(self.users, "users.json")
                messagebox.showinfo("Success", "Bank account updated successfully!")
            else:
                messagebox.showerror("Error", "Please enter a valid account number")
    
    # Method to create employee account
    def create_employee_account(self):
        create_window = ctk.CTkToplevel(self.root)
        create_window.title("Create Employee Account")
        create_window.geometry("400x300")

        ctk.CTkLabel(create_window, text="Create Employee Account", font=("Arial", 20, "bold")).pack(pady=(20,10))

        ctk.CTkLabel(create_window, text="Username:", font=("Arial", 14)).pack(pady=5)
        new_username = ctk.CTkEntry(create_window, width=300)
        new_username.pack(pady=5)

        ctk.CTkLabel(create_window, text="Password:", font=("Arial", 14)).pack(pady=5)
        new_password = ctk.CTkEntry(create_window, show="*", width=300)
        new_password.pack(pady=5)

        def submit():
            username = new_username.get()
            password = new_password.get()
            if username and password:
                if username not in self.users:
                    self.users[username] = {
                        'password': password,
                        'role': 'employee',
                        'business': self.current_user.business
                    }
                    self.save_data(self.users, "users.json")
                    messagebox.showinfo("Success", f"Employee account '{username}' created!")
                    create_window.destroy()
                else:
                    messagebox.showerror("Error", "Username already exists")
            else:
                messagebox.showerror("Error", "Please fill all fields")

        ctk.CTkButton(create_window, text="Create Account", command=submit, 
                    fg_color="#4CAF50", hover_color="#45a049", width=250).pack(pady=20)

    # Method to change employee password
    def change_employee_password(self):
        change_window = ctk.CTkToplevel(self.root)
        change_window.title("Change Employee Password")
        change_window.geometry("400x300")

        ctk.CTkLabel(change_window, text="Change Employee Password", font=("Arial", 20, "bold")).pack(pady=(20,10))

        ctk.CTkLabel(change_window, text="Employee Username:", font=("Arial", 14)).pack(pady=5)
        employee_username = ctk.CTkEntry(change_window, width=300)
        employee_username.pack(pady=5)

        ctk.CTkLabel(change_window, text="New Password:", font=("Arial", 14)).pack(pady=5)
        new_password = ctk.CTkEntry(change_window, show="*", width=300)
        new_password.pack(pady=5)

        def submit():
            username = employee_username.get()
            password = new_password.get()
            if username in self.users and self.users[username]['role'] == 'employee' and self.users[username]['business'] == self.current_user.business:
                self.users[username]['password'] = password
                self.save_data(self.users, "users.json")
                messagebox.showinfo("Success", f"Password changed for employee '{username}'")
                change_window.destroy()
            else:
                messagebox.showerror("Error", "Invalid employee username")

        ctk.CTkButton(change_window, text="Change Password", command=submit, 
                    fg_color="#2196F3", hover_color="#1976D2", width=250).pack(pady=20)

    # Method to set up the invoice frame
    def setup_invoice_frame(self):
        # Add customer name entry
        ctk.CTkLabel(self.invoice_frame, text="Customer Name:").pack(pady=(10, 5))
        self.customer_name_entry = ctk.CTkEntry(self.invoice_frame, width=300)
        self.customer_name_entry.pack(pady=5)

        # Add customer email entry
        ctk.CTkLabel(self.invoice_frame, text="Customer Email:").pack(pady=(10, 5))
        self.customer_email_entry = ctk.CTkEntry(self.invoice_frame, width=300)
        self.customer_email_entry.pack(pady=5)

        # Add invoice text area
        self.invoice_text = ctk.CTkTextbox(self.invoice_frame, width=500, height=400)
        self.invoice_text.pack(pady=10)

        # Buttons for saving invoice and going back
        ctk.CTkButton(self.invoice_frame, text="Save Invoice", command=self.save_invoice).pack(pady=10)
        ctk.CTkButton(self.invoice_frame, text="Back to Products", command=lambda: self.show_frame(self.product_frame)).pack(pady=10)

    # Method to set up the history frame
    def setup_history_frame(self):
        self.history_frame.grid_columnconfigure(0, weight=1)
        self.history_frame.grid_rowconfigure(1, weight=1)

        # Create filter frame
        filter_frame = ctk.CTkFrame(self.history_frame)
        filter_frame.grid(row=0, column=0, padx=10, pady=10, sticky="ew")

        # Add user filter
        ctk.CTkLabel(filter_frame, text="Filter by User:").pack(side="left", padx=5)
        self.user_filter = ctk.CTkComboBox(filter_frame, values=["All Users"])
        self.user_filter.pack(side="left", padx=5)
        self.user_filter.set("All Users")

        # Add price filter
        ctk.CTkLabel(filter_frame, text="Filter by Price:").pack(side="left", padx=5)
        self.price_filter = ctk.CTkEntry(filter_frame, width=100)
        self.price_filter.pack(side="left", padx=5)

        # Add apply filters button
        ctk.CTkButton(filter_frame, text="Apply Filters", command=self.apply_filters).pack(side="left", padx=5)

        # Add history text area
        self.history_tree = ctk.CTkTextbox(self.history_frame, width=500, height=400)
        self.history_tree.grid(row=1, column=0, padx=10, pady=10, sticky="nsew")

        # Add back button
        ctk.CTkButton(self.history_frame, text="Back to Products", command=lambda: self.show_frame(self.product_frame)).grid(row=2, column=0, pady=10)
    
    # Method to show a specific frame
    def show_frame(self, frame):
        for f in (self.login_frame, self.signup_frame, self.product_frame, self.invoice_frame, self.history_frame):
            f.pack_forget()
        frame.pack(expand=True, fill="both")

    # Method to handle login
    def login(self):
        username = self.username_entry.get()
        password = self.password_entry.get()
        if username in self.users and self.users[username]['password'] == password:
            user_data = self.users[username]
            self.current_user = User(username, password, user_data['role'], user_data['business'])
            self.show_frame(self.product_frame)
            self.load_products()
            
            # Show/hide Store tab based on user role
            if self.current_user.role == "owner":
                # Check if Store tab exists before adding
                if "Store" not in self.tab_view._tab_dict:
                    self.tab_view.add("Store")
                    self.setup_store_tab()
            else:
                # If user is not owner, try to delete Store tab if it exists
                try:
                    self.tab_view.delete("Store")
                except:
                    pass  # If Store tab doesn't exist, just move on
            
            messagebox.showinfo("Success", f"Welcome, {username}!")
        else:
            messagebox.showerror("Error", "Invalid username or password")
    
    # Method to handle signup
    def signup(self):
        username = self.new_username_entry.get()
        password = self.new_password_entry.get()
        business = self.business_entry.get()
        if username and password and business:
            if username not in self.users:
                self.users[username] = {
                    'password': password,
                    'role': 'owner',  # Always set role to 'owner'
                    'business': business
                }
                self.save_data(self.users, "users.json")
                
                # Initialise empty product list for the new business
                if business not in self.products:
                    self.products[business] = []
                    self.save_data(self.products, "products.json")
                
                messagebox.showinfo("Success", "Owner account created! You can now log in.")
                self.show_frame(self.login_frame)
            else:
                messagebox.showerror("Error", "Username already exists")
        else:
            messagebox.showerror("Error", "Please fill all fields")
    
    # Method to add a new product
    def add_product(self):
        if self.current_user.role != "owner":
            messagebox.showerror("Error", "Only owners can add products.")
            return

        name = self.product_name_entry.get()
        price = self.product_price_entry.get()
        try:
            price = float(price)
            if price < 0:
                raise ValueError
            if self.current_user.business not in self.products:
                self.products[self.current_user.business] = []
            self.products[self.current_user.business].append({
                'name': name,
                'price': price,
                'added_by': self.current_user.username
            })
            self.save_data(self.products, "products.json")
            self.load_products()
            self.product_name_entry.delete(0, ctk.END)
            self.product_price_entry.delete(0, ctk.END)
            messagebox.showinfo("Success", f"Product '{name}' added successfully.")
        except ValueError:
            messagebox.showerror("Error", "Invalid price. Please enter a positive number.")

    # Method to load products for the current business
    def load_products(self):
        if self.current_user.business in self.products:
            product_names = [product['name'] for product in self.products[self.current_user.business]]
            self.product_listbox.configure(values=product_names)
            if product_names:
                self.product_listbox.set(product_names[0])
                self.update_price_display(product_names[0])
        else:
            # If no products exist for this business, clear the product list
            self.product_listbox.configure(values=[])
            self.product_listbox.set("")
            self.price_display.configure(text="Price: $0.00")

    # Method to update the price display when a product is selected
    def update_price_display(self, selected_product_name):
        for product in self.products.get(self.current_user.business, []):
            if product['name'] == selected_product_name:
                self.price_display.configure(text=f"Price: ${product['price']:.2f}")
                break
    
    # Method to add a product to the invoice
    def add_to_invoice(self):
        selected_product_name = self.product_listbox.get()
        quantity = self.quantity_entry.get()
        
        try:
            quantity = int(quantity)
            if quantity <= 0:
                raise ValueError
        except ValueError:
            messagebox.showerror("Error", "Please enter a valid quantity")
            return
        
        for product in self.products[self.current_user.business]:
            if product['name'] == selected_product_name:
                product_with_quantity = product.copy()
                product_with_quantity['quantity'] = quantity
                self.selected_products.append(product_with_quantity)
                self.update_selected_products_text()
                break
        
        self.quantity_entry.delete(0, ctk.END)
        self.selected_products_text.configure(state='disabled')

    # Method to update the text showing selected products
    def update_selected_products_text(self):
        self.selected_products_text.delete('1.0', ctk.END)
        for product in self.selected_products:
            self.selected_products_text.insert(ctk.END, f"{product['name']} - ${product['price']:.2f} x {product['quantity']} = ${product['price'] * product['quantity']:.2f}\n")

    # Method to show the invoice                       
    def show_invoice(self):
        if not self.selected_products:
            messagebox.showerror("Error", "No products selected for the invoice")
            return

        self.customer_name_entry.delete(0, ctk.END)
        self.customer_email_entry.delete(0, ctk.END)

        total = sum(product['price'] * product['quantity'] for product in self.selected_products)
        invoice_text = f"Invoice for {self.current_user.business}\n"
        invoice_text += f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        for product in self.selected_products:
            invoice_text += f"{product['name']}: ${product['price']:.2f} x {product['quantity']} = ${product['price'] * product['quantity']:.2f}\n"
        invoice_text += f"\nTotal: ${total:.2f}"

        self.invoice_text.delete('1.0', ctk.END)
        self.invoice_text.insert(ctk.END, invoice_text)
        self.show_frame(self.invoice_frame)

    # Method to save the invoice
    def save_invoice(self):
        customer_name = self.customer_name_entry.get().strip()
        customer_email = self.customer_email_entry.get().strip()

        if not customer_name or not customer_email:
            messagebox.showerror("Error", "Please enter both customer name and email.")
            return

        # Create a new Word document
        doc = Document()
        
        # Add invoice details to the document
        doc.add_heading(self.current_user.business, 0)
        
        doc.add_paragraph(f"Customer: {customer_name}")
        doc.add_paragraph(f"Customer Email: {customer_email}")
        doc.add_paragraph(f"Employee: {self.current_user.username}")
        doc.add_paragraph(f"Date and Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Bank Account: {self.users[self.current_user.username].get('bank_account', 'Not provided')}")
        
        # Create a table for the products
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Product'
        hdr_cells[1].text = 'Quantity'
        hdr_cells[2].text = 'Price'
        hdr_cells[3].text = 'Total'
        
        # Add products to the table
        for product in self.selected_products:
            row_cells = table.add_row().cells
            row_cells[0].text = product['name']
            row_cells[1].text = str(product['quantity'])
            row_cells[2].text = f"${product['price']:.2f}"
            row_cells[3].text = f"${product['price'] * product['quantity']:.2f}"
        
        # Add total to the document
        total = sum(product['price'] * product['quantity'] for product in self.selected_products)
        doc.add_paragraph(f"Total: ${total:.2f}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Save the document
        filename = f"invoice_{self.current_user.business}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        doc.save(filename)

        # Add to purchase history
        if self.current_user.business not in self.purchase_history:
            self.purchase_history[self.current_user.business] = []
        
        self.purchase_history[self.current_user.business].append({
            'date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'total': total,
            'created_by': self.current_user.username,
            'customer_name': customer_name,
            'customer_email': customer_email
        })
        self.save_data(self.purchase_history, "purchase_history.json")
        
        messagebox.showinfo("Success", f"Invoice saved to {filename}")
        self.selected_products = []
        self.update_selected_products_text()
        self.customer_name_entry.delete(0, ctk.END)
        self.customer_email_entry.delete(0, ctk.END)

    # Method to show purchase history
    def show_history(self):
        self.history_tree.delete('1.0', ctk.END)
        users = set()
        if self.current_user.business in self.purchase_history:
            for purchase in self.purchase_history[self.current_user.business]:
                if 'created_by' in purchase:
                    users.add(purchase['created_by'])
                else:
                    users.add("Unknown")
        self.user_filter.configure(values=["All Users"] + list(users))
        self.apply_filters()
        self.show_frame(self.history_frame)
    
    # Method to apply filters to the purchase history
    def apply_filters(self):
        self.history_tree.delete('1.0', ctk.END)
        selected_user = self.user_filter.get()
        price_filter = self.price_filter.get()

        try:
            price_filter = float(price_filter) if price_filter else None
        except ValueError:
            messagebox.showerror("Error", "Invalid price filter. Please enter a valid number.")
            return

        if self.current_user.business in self.purchase_history:
            for purchase in self.purchase_history[self.current_user.business]:
                created_by = purchase.get('created_by', "Unknown")
                total = purchase['total']

                if (selected_user == "All Users" or selected_user == created_by) and \
                   (price_filter is None or total == price_filter):
                    self.history_tree.insert(ctk.END, f"Date: {purchase['date']}\n")
                    self.history_tree.insert(ctk.END, f"Created by: {created_by}\n")
                    self.history_tree.insert(ctk.END, f"Customer: {purchase.get('customer_name', 'N/A')}\n")
                    self.history_tree.insert(ctk.END, f"Email: {purchase.get('customer_email', 'N/A')}\n")
                    self.history_tree.insert(ctk.END, f"Total: ${total:.2f}\n")
                    self.history_tree.insert(ctk.END, "-" * 40 + "\n")

    def logout(self):
        self.current_user = None
        self.show_frame(self.login_frame)

    def run(self):
        self.root.mainloop()

if __name__ == "__main__":
    app = InvoiceMaker()
    app.run()